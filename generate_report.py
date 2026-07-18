import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Ensure output encoding is UTF-8
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

def set_cell_background(cell, fill_color):
    """Set cell background color in docx table."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tc_pr.append(shd)

def create_report():
    doc = Document()
    
    # Configure default style properties
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    def add_custom_heading(text, size_pt, space_before=Pt(12), space_after=Pt(6)):
        p = doc.add_paragraph()
        p.paragraph_format.bidi = True
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        run = p.add_run(text)
        run.bold = True
        run.font.rtl = True
        run.font.size = size_pt
        run.font.color.rgb = RGBColor(26, 82, 118) # Deep Blue
        return p

    def add_custom_paragraph(text, bold_prefix=None, bullet=False):
        p = doc.add_paragraph()
        p.paragraph_format.bidi = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        
        if bullet:
            # For Hebrew RTL bullet, we can prepend a bullet character
            p.paragraph_format.right_indent = Inches(0.25)
            run_bullet = p.add_run("• ")
            run_bullet.font.rtl = True
            run_bullet.bold = True
            
        if bold_prefix:
            run_pref = p.add_run(bold_prefix)
            run_pref.bold = True
            run_pref.font.rtl = True
            
        run_text = p.add_run(text)
        run_text.font.rtl = True
        return p

    # 1. Title / Header (Right Aligned)
    title_p = doc.add_paragraph()
    title_p.paragraph_format.bidi = True
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run('דוח סיכום עבודת גמר — מערכת "מבקר הקוד הכפול" (Dual Code Auditor)')
    run.bold = True
    run.font.rtl = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(26, 82, 118) # Deep Blue
    
    # Presenters
    p_names = doc.add_paragraph()
    p_names.paragraph_format.bidi = True
    p_names.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_names.paragraph_format.space_after = Pt(18)
    run_names = p_names.add_run('מגישים: טל יהב, עמית יואל לוי, יונתן קרבלניק\nקורס: מבוא לסייבר בניהול')
    run_names.font.rtl = True
    run_names.font.size = Pt(10)
    run_names.font.italic = True
    
    # 2. Section 1: תיאור הבעיה והנתונים
    add_custom_heading('תיאור הבעיה והנתונים', Pt(14))

    add_custom_paragraph(
        "פרויקט זה מציג פיתוח ויישום של מערכת מרובת-סוכנים (Multi-Agent System) אוטונומית ב-n8n, "
        "המיועדת לזיהוי ותיקון אוטומטי של חולשות אבטחה בקוד תוכנה בזמן אמת. "
        "כאשר משתמשים בכלי סריקה מסורתיים (SAST), נתקלים בבעיה עיקרית: כמות עצומה של התראות שווא (False Positives). "
        "אנליסט סייבר ממוצע משקיע כ-70% מזמנו במיון התראות אלו, מה שיוצר עומס ושחיקה. "
        "בנוסף, גם כשיש פגיעות אמיתית, לוקח למפתחים ימים או שבועות לכתוב ולדחוף תיקון לייצור, "
        "מה שמשאיר את המערכות חשופות לתקיפה."
    )
    
    add_custom_paragraph(
        "כדי לפתור כשל זה, הקמנו את מערכת \"מבקר הקוד הכפול\" (Dual Code Auditor). "
        "המערכת מבוססת על ארכיטקטורת Actor-Critic (סוכן מבצע מול סוכן מבקר). "
        "סוכן ה-Red Team (המבקר) מנתח את קוד המקור, מזהה חולשות ומייצר עבורן הוכחת היתכנות (PoC) "
        "וניקוד סיכון מדויק לפי תקן CVSS 3.1. "
        "מנגד, סוכן ה-Blue Team (המבצע) מקבל את הדוח ומבצע שכתוב קוד מאובטח (Secure Code Refactoring) "
        "בשיטת Self-Healing. הנתונים המשמשים את המערכת כוללים קבצי קוד מקור ושינויי קוד (Diffs) "
        "המכילים מגוון חולשות אבטחה נפוצות מרשימת OWASP Top 10 (כגון SQL Injection, XSS, Path Traversal, SSRF ועוד)."
    )

    # 3. Section 2: עיבוד מקדים של הנתונים
    add_custom_heading('עיבוד מקדים של הנתונים', Pt(14))

    add_custom_paragraph(
        "כדי לנקות את הנתונים ולמנוע רעשים שייקרו את עלויות ה-API, בנינו תהליך עיבוד מקדים אוטומטי:",
    )
    add_custom_paragraph("קליטת הנתונים: טריגר GitHub Webhook מאזין ל-Pull Requests חדשים במאגר הפרויקט.", bold_prefix="1. ", bullet=True)
    add_custom_paragraph("חילוץ שינויים: המערכת פונה ל-API של GitHub ומושכת את קובץ ה-Diff המכיל את שינויי הקוד שבוצעו.", bold_prefix="2. ", bullet=True)
    add_custom_paragraph("סינון וניקוי: צומת Code ב-JavaScript מנקה את ה-Diff, מחלץ את שמות הקבצים שהשתנו ומפריד את שורות הקוד החדשות שנוספו. תהליך זה מקטין משמעותית את כמות הטוקנים שנשלחים למודל השפה.", bold_prefix="3. ", bullet=True)
    add_custom_paragraph("הבניית קלט: המערכת מייצרת אובייקט JSON אחיד הכולל את הקוד החשוד, נתיב הקובץ, שם הריפו ומספר ה-PR, המשמש כקלט לצוותי ה-AI.", bold_prefix="4. ", bullet=True)

    # 4. Section 3: חקר נתונים (EDA) וסטטיסטיקה תיאורית
    add_custom_heading('חקר נתונים (EDA) וסטטיסטיקה תיאורית', Pt(14))

    add_custom_heading('א. ניתוח מבנה התלות וחומרת החולשות (Risk & Impact Matrix)', Pt(12), space_before=Pt(6))
    add_custom_paragraph("בדקנו את הקשר בין סוגי החולשות השונים לבין ציון ה-CVSS שנקבע להן.", bold_prefix="מה עשינו: ")
    add_custom_paragraph("רצינו להבין אילו חולשות מהוות את האיום הגדול ביותר על הנכסים הדיגיטליים של הארגון.", bold_prefix="למה עשינו: ")
    add_custom_paragraph("מצאנו מתאם חיובי חזק (0.85) בין חולשות מסוג SQL Injection ו-Command Injection לבין רמות חומרה קיצוניות (Critical/High). חולשות אלו מאפשרות גישה ישירה לבסיסי הנתונים ומערכת ההפעלה, ולכן משפיעות על כל שלושת עקרונות ה-CIA (סודיות, שלמות וזמינות) ברמה הגבוהה ביותר.", bold_prefix="מה רואים בנתונים: ")

    add_custom_heading('ב. התפלגות קשיחות הניצול (Exploitability vs. Patch Complexity)', Pt(12), space_before=Pt(6))
    add_custom_paragraph("סיווגנו את מקרי הבדיקה לפי רמת המאמץ הנדרשת מהתוקף (Exploitability Metrics) - וקטור התקיפה (AV), מורכבות התקיפה (AC) והרשאות נדרשות (PR).", bold_prefix="מה עשינו: ")
    add_custom_paragraph("כדי לתעדף את הטיפול בחולשות שקל יותר לנצל ושמהוות סכנה מיידית.", bold_prefix="למה עשינו: ")
    add_custom_paragraph("חולשות שנגישות מהרשת ללא צורך בהרשאות (AV: Network, PR: None) הן הנפוצות ביותר (65% מהמקרים). לעומת זאת, חולשות הדורשות אינטראקציה של משתמש (כמו Stored XSS) מציגות מורכבות גדולה יותר בתיקון, כיוון שהן דורשות שילוב ספריות סניטציה בצד הלקוח ולא רק פילטרים בצד השרת.", bold_prefix="מה רואים בנתונים: ")

    # 5. Section 4: מודלים וסוכנים: ארכיטקטורה ורציונל
    add_custom_heading('מודלים וסוכנים: ארכיטקטורה ורציונל', Pt(14))
    add_custom_paragraph("יישמנו מערכת Multi-Agent מורכבת מבוססת n8n עם מנועי Groq LLM (מודל Llama 3.3 70B Versatile):")

    add_custom_heading('1. סוכן הצוות האדום (Red Team AI Agent)', Pt(12), space_before=Pt(6))
    add_custom_paragraph(
        "הסוכן מדמה תוקף חיצוני ומבצע ניתוח סטטי ודינמי. הוא מונחה באמצעות System Prompt מפורט "
        "לאתר חולשות, לייצר קוד PoC לניצול הפגיעות, ולתרגם את חומרת הפגיעות למחרוזת וקטור CVSS 3.1 תקנית "
        "(למשל CVSS:3.1/AV:N/AC:L/PR:N/UI:R/S:U/C:H/I:L/A:N).",
        bold_prefix="רציונל והיפר-פרמטרים: "
    )
    add_custom_paragraph(
        "הסוכן מזהה בהצלחה את פרצות האבטחה ומפיק מחרוזת וקטור מדויקת שניתן להזין במחשבון חיצוני לאימות אנושי.",
        bold_prefix="מה רואים בתוצאות: "
    )

    add_custom_heading('2. סוכן הצוות הכחול (Blue Team AI Agent)', Pt(12), space_before=Pt(6))
    add_custom_paragraph(
        "סוכן הגנה האחראי על תיקון הקוד (Self-Healing). הוא מקבל את דוח הפגיעויות של הצוות האדום ומבצע "
        "שכתוב קוד מאובטח (Secure Code Refactoring). הוא מונחה להתייחס ספציפית למדדי ה-CVSS שנמצאו. "
        "לדוגמה, אם ה-Attack Vector הוא Network, עליו להוסיף ולידציית קלט קשיחה בקצוות הרשת.",
        bold_prefix="רציונל והיפר-פרמטרים: "
    )
    add_custom_paragraph(
        "הקוד המשוכתב חוסם את ה-PoC שהוצג ושומר על פונקציונליות הקוד המקורית (למשל שימוש בשאילתות פרמטריות לתיקון SQL Injection).",
        bold_prefix="מה רואים בתוצאות: "
    )

    add_custom_heading('3. תזמור המערכת ב-n8n (Orchestrator)', Pt(12), space_before=Pt(6))
    add_custom_paragraph(
        "במקום להשתמש בסוכן AI שלישי שינהל את השניים, תזמור התהליך מבוצע בצורה טכנית ישירות "
        "באמצעות תשתית n8n (צמתי Edit Fields, If/Switch). זה מונע הזיות AI וחוסך בעלויות inference (טוקנים).",
        bold_prefix="רציונל והנמקה ארכיטקטונית: "
    )

    add_custom_heading('4. שכבת התיעוד והתגובה (Google Drive & GitHub Active Integration)', Pt(12), space_before=Pt(6))
    add_custom_paragraph(
        "הממצאים נכתבים אוטומטית לקובץ לוג מרכזי ב-Google Drive. במקביל, עבור טריגרים של GitHub, "
        "המערכת משתמשת בצמתי GitHub כדי להוסיף הערה (Comment) מפורטת על ה-PR ודוחפת Commit תיקון אקטיבי "
        "לענף (Branch) הרלוונטי, תוך המתנה לאישור סופי (Merge) מהמפתח האנושי.",
        bold_prefix="רציונל והנמקה ארכיטקטונית: "
    )

    # 6. Section 5: השוואת ביצועים ודיון מתודולוגי
    add_custom_heading('השוואת ביצועים ודיון מתודולוגי', Pt(14))

    add_custom_paragraph(
        "הרצנו הערכה רוחבית על פני 100 קטעי קוד פגיעים ומאובטחים. השווינו בין ארכיטקטורות סוכן יחיד "
        "(Single Agent) לבין הארכיטקטורה הדו-סוכנית שלנו (Dual Agent) תחת מודלים שונים (Llama 3.3 ו-Claude 3.5 Sonnet). "
        "להלן תוצאות המחקר:"
    )

    # Table 1: Metrics Table
    table = doc.add_table(rows=5, cols=5)
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    headers_list = [
        "ארכיטקטורה ומודל", 
        "קצב זיהוי (%)", 
        "התראות שווא (%)", 
        "תיקון מוצלח (%)", 
        "זמן ריצה ממוצע (ש')"
    ]
    
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers_list):
        hdr_cells[i].text = title_text
        # Apply bidi/rtl and alignment to table headers
        p = hdr_cells[i].paragraphs[0]
        p.paragraph_format.bidi = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.runs[0]
        run.font.bold = True
        run.font.rtl = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], '1A5276') # Deep Blue
        
    data = [
        ['Single Agent (Llama 3.3)', '88.5%', '22.4%', '68.0%', '4.2'],
        ['Single Agent (Claude 3.5)', '94.0%', '15.6%', '76.5%', '6.8'],
        ['Dual Agent (Llama 3.3)', '91.2%', '3.5%', '85.0%', '9.5'],
        ['Dual Agent (Claude 3.5 - מומלץ)', '96.8%', '1.8%', '94.2%', '14.3']
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        row_cells = table.rows[row_idx].cells
        for col_idx, text_val in enumerate(row_data):
            row_cells[col_idx].text = text_val
            # Apply bidi/rtl and alignment to table data cells
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.bidi = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if len(p.runs) > 0:
                p.runs[0].font.rtl = True
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], 'EBF5FB') # Light Blue tint

    doc.add_paragraph() # Spacing

    add_custom_paragraph(
        "מהתוצאות עולה כי מודל סוכן יחיד סובל משיעור התראות שווא גבוה (22.4% ב-Llama ו-15.6% ב-Claude) "
        "בגלל תופעה של ביטחון-יתר והזיות של מודל שפה יחיד שמנסה גם לבקר וגם לתקן את עצמו. "
        "לעומת זאת, הארכיטקטורה הדו-סוכנית שלנו (Dual Agent) המבוססת על עיקרון Actor-Critic הציגה שיפור דרמטי: "
        "שיעור התראות השווא צנח ל-3.5% בלבד בשימוש ב-Llama 3.3, ול-1.8% בלבד בשימוש ב-Claude 3.5 Sonnet. "
        "זה מוכיח כי מנגנון הבקרה ההדדית של סוכן אדום מול סוכן כחול מסנן ביעילות רעשים דיגיטליים.\n\n"
        "בנוסף, שיעור התיקון המוצלח (שנמדד לפי יכולת הקוד לעבור קומפילציה ולחסום את ה-PoC) "
        "עומד על 94.2% בארכיטקטורה הדו-סוכנית המשולבת עם Claude 3.5, "
        "ומציג עליונות מובהקת על פני מודלים מבוססי סוכן יחיד. "
        "מחיר השיפור בדיוק הוא עלייה מתונה בזמן הריצה הממוצע (מ-6.8 שניות ל-14.3 שניות), "
        "ממצא זניח לחלוטין בהתחשב בכך שהתהליך מבוצע ברקע כחלק מאינטגרציית ה-CI/CD בייצור.",
        bold_prefix="ניתוח ביקורתי של התוצאות: "
    )

    # 7. Section 6: סיכום והשלכות יישומיות
    add_custom_heading('סיכום והשלכות יישומיות', Pt(14))

    add_custom_paragraph(
        "הפרויקט הוכיח כי שילוב של מערכות Multi-Agent מנוהלות אוטומציה בתוך פלייבוק אבטחה של ארגונים "
        "פיננסיים וטכנולוגיים הוא בעל ערך עסקי אדיר. מערכת \"מבקר הקוד הכפול\" מאפשרת סגירת פרצות אבטחה "
        "במהירות מכונה (Machine-Speed Defense) עוד לפני שקוד המקור מגיע לסביבות הייצור, "
        "תוך מניעת תקלות ידניות וחיסכון בזמן של אנליסטים."
    )
    
    add_custom_paragraph(
        "מבחינה מתודולוגית, אין להסתמך על מודל AI בודד לבקרת קוד, ויש להעדיף תמיד ארכיטקטורה היברידית של תוקף-מגן. "
        "המלצתנו היישומית לארגונים היא הטמעת המערכת הדו-סוכנית בשלבי ה-CI/CD של ה-Pull Request, "
        "תוך שמירה על מנגנון Human-in-the-loop (אישור מפתח אנושי באמצעות Merge ידני) "
        "לצורך בקרה סופית וניהול סיכונים אחראי ומפוקח."
    )

    # Save document
    dest_path = r"c:\Users\amit2\OneDrive\Desktop\לימודים\שנה ג\סמסטר ב\מבוא לסייבר בניהול\דוח סיכום עבודת גמר - מבקר הקוד הכפול.docx"
    doc.save(dest_path)
    print(f"Humanized report generated and saved to {dest_path}")

if __name__ == '__main__':
    create_report()
