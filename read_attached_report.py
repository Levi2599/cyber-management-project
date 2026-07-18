from docx import Document

path = r"C:\Users\amit2\AppData\Local\Temp\דוח סיכום עבודת גמר.docx"
output_file = r"C:\tmp\attached_report.txt"

with open(output_file, 'w', encoding='utf-8') as out:
    try:
        doc = Document(path)
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                out.write(text + "\n")
        for i, table in enumerate(doc.tables, 1):
            out.write(f"\n[TABLE {i}]\n")
            for row in table.rows:
                out.write(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells) + "\n")
        print("Success. Saved to C:\\tmp\\attached_report.txt")
    except Exception as e:
        print(f"Error: {e}")
