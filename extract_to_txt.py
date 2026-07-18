from pathlib import Path
from docx import Document
from pypdf import PdfReader
import json

root = Path(r"C:\Users\amit2\OneDrive\Desktop\לימודים\שנה ג\סמסטר ב\מבוא לסייבר בניהול")
output_file = Path(r"C:\tmp\all_docs.txt")

with open(output_file, 'w', encoding='utf-8') as out:
    # 1. Read PRD (v2)
    prd_path = root / "PRD - מבקר הקוד הכפול (v2).docx"
    out.write(f"\n=========================================\nFILE: {prd_path.name}\n=========================================\n")
    try:
        doc = Document(prd_path)
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                out.write(text + "\n")
        for i, table in enumerate(doc.tables, 1):
            out.write(f"\n[TABLE {i}]\n")
            for row in table.rows:
                out.write(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells) + "\n")
    except Exception as e:
        out.write(f"Error reading {prd_path.name}: {e}\n")

    # 2. Read Technical plan
    tech_path = root / "תוכנית עבודה טכנית_ הקמת מערכת 'מבקר הקוד הכפול' ב-n8n.docx"
    out.write(f"\n=========================================\nFILE: {tech_path.name}\n=========================================\n")
    try:
        doc = Document(tech_path)
        for p in doc.paragraphs:
            text = p.text.strip()
            if text:
                out.write(text + "\n")
        for i, table in enumerate(doc.tables, 1):
            out.write(f"\n[TABLE {i}]\n")
            for row in table.rows:
                out.write(" | ".join(cell.text.replace("\n", " / ").strip() for cell in row.cells) + "\n")
    except Exception as e:
        out.write(f"Error reading {tech_path.name}: {e}\n")

    # 3. Read PDF
    pdf_path = root / "פרויקט גמר קורס סייבר בעידן ה.pdf"
    out.write(f"\n=========================================\nFILE: {pdf_path.name}\n=========================================\n")
    try:
        reader = PdfReader(str(pdf_path))
        for i, page in enumerate(reader.pages, 1):
            out.write(f"\n[PAGE {i}]\n")
            out.write((page.extract_text() or "").strip() + "\n")
    except Exception as e:
        out.write(f"Error reading {pdf_path.name}: {e}\n")

print("Done. Saved to C:\\tmp\\all_docs.txt")
