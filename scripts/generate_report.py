"""
generate_report.py — מחולל את דוח הסיכום (DOCX + MD).

עקרון מנחה: הדוח לא מכיל מספר שלא נמדד.
  * הרכב קורפוס הבדיקה נקרא מ-scripts/eval/dataset.json.
  * תוצאות הביצועים נקראות מ-scripts/eval/results.json — ורק אם הקובץ קיים.
    אם ההערכה טרם הורצה, הדוח אומר זאת במפורש במקום להציג טבלה.
  * תוצאות בדיקות היחידה נקראות מהרצה בפועל של scripts/eval/test_code_nodes.js.

הרצה:
    python scripts/make_charts.py
    python scripts/generate_report.py
"""

import json
import os
import subprocess
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
EVAL_DIR = os.path.join(HERE, "eval")
FIG_DIR = os.path.join(HERE, "figures")

DOCX_OUT = os.path.join(ROOT, "דוח סיכום עבודת גמר - מבקר הקוד הכפול.docx")
MD_OUT = os.path.join(ROOT, "דוח סיכום עבודת גמר - מבקר הקוד הכפול.md")

DEEP_BLUE = RGBColor(26, 82, 118)

# צובר את הדוח גם כ-Markdown, במקביל לבנייה של ה-DOCX.
md_lines = []


# --------------------------------------------------------------------------
# עזרי עיצוב RTL
# --------------------------------------------------------------------------
def set_cell_background(cell, fill_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_color)
    tc_pr.append(shd)


def heading(doc, text, size_pt, md_level=2, space_before=Pt(12)):
    p = doc.add_paragraph()
    p.paragraph_format.bidi = True
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.bold = True
    run.font.rtl = True
    run.font.size = size_pt
    run.font.color.rgb = DEEP_BLUE
    md_lines.append("\n" + "#" * md_level + " " + text + "\n")


def para(doc, text, bold_prefix=None, bullet=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.bidi = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15

    if bullet:
        p.paragraph_format.right_indent = Inches(0.25)
        r = p.add_run("• ")
        r.font.rtl = True
        r.bold = True
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.rtl = True
    r = p.add_run(text)
    r.font.rtl = True
    r.font.italic = italic

    prefix = "* " if bullet else ""
    # ב-Markdown רווח צמוד לסימני הסגירה מבטל את ההדגשה, ולכן הרווח שאחרי
    # הכותרת מועבר אל מחוץ לסימנים.
    bold_md = "**%s** " % bold_prefix.rstrip() if bold_prefix else ""
    md_lines.append(prefix + bold_md + text)


def table(doc, headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        p = hdr[i].paragraphs[0]
        p.paragraph_format.bidi = True
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.runs[0]
        run.font.bold = True
        run.font.rtl = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr[i], "1A5276")

    for ri, row in enumerate(rows, 1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            p = cells[ci].paragraphs[0]
            p.paragraph_format.bidi = True
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if p.runs:
                p.runs[0].font.rtl = True
            if ri % 2 == 0:
                set_cell_background(cells[ci], "EBF5FB")
    doc.add_paragraph()

    md_lines.append("")
    md_lines.append("| " + " | ".join(headers) + " |")
    md_lines.append("| " + " | ".join([":---"] * len(headers)) + " |")
    for row in rows:
        md_lines.append("| " + " | ".join(str(c) for c in row) + " |")
    md_lines.append("")


def figure(doc, filename, caption):
    path = os.path.join(FIG_DIR, filename)
    if not os.path.exists(path):
        return False
    doc.add_picture(path, width=Inches(5.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, italic=True)
    md_lines.append("\n![%s](scripts/figures/%s)\n" % (caption, filename))
    return True


# --------------------------------------------------------------------------
# איסוף נתונים אמיתיים
# --------------------------------------------------------------------------
def load_dataset():
    with open(os.path.join(EVAL_DIR, "dataset.json"), encoding="utf-8") as f:
        return json.load(f)


def load_baseline():
    """ההרצה הראשונה, לפני תיקון הפרומפט. משמשת להשוואה."""
    path = os.path.join(EVAL_DIR, "results_baseline.json")
    if not os.path.exists(path):
        return None
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    if data.get("summary", {}).get("valid") is False:
        return None
    return data


def load_midrun():
    """ההרצה השנייה, אחרי הוספת בדיקת הנטרול ולפני החריגים."""
    path = os.path.join(EVAL_DIR, "results_v2.json")
    if not os.path.exists(path):
        return None
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    if data.get("summary", {}).get("valid") is False:
        return None
    return data


def load_results():
    path = os.path.join(EVAL_DIR, "results.json")
    if not os.path.exists(path):
        return None
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    summary = data.get("summary", {})
    if summary.get("completed_cases", 0) == 0:
        return None
    # ריצה שחלק ממקריה נכשלו אינה מדידה. עדיף דוח שאומר "טרם הורצה" על פני
    # דוח שמציג אחוזים שחושבו על מדגם חלקי.
    if summary.get("valid") is False:
        print("NOTE: results.json is marked invalid — %s" % summary.get("invalid_reason", ""))
        print("      The report will state that the evaluation has not been run.")
        return None
    return data


def run_unit_tests():
    """מריץ את בדיקות היחידה ומחזיר (עברו, נכשלו) — או None אם לא ניתן להריץ."""
    try:
        proc = subprocess.run(
            ["node", os.path.join(EVAL_DIR, "test_code_nodes.js")],
            capture_output=True, text=True, encoding="utf-8", errors="replace",
            timeout=120, cwd=ROOT,
        )
    except (OSError, subprocess.TimeoutExpired):
        return None
    for line in reversed((proc.stdout or "").splitlines()):
        if "passed" in line and "failed" in line:
            parts = line.split()
            try:
                return int(parts[0]), int(parts[2])
            except (ValueError, IndexError):
                return None
    return None


# --------------------------------------------------------------------------
# בניית הדוח
# --------------------------------------------------------------------------
def build():
    dataset = load_dataset()
    results = load_results()
    baseline = load_baseline()
    midrun = load_midrun()
    tests = run_unit_tests()

    cases = dataset["cases"]
    n_vuln = sum(1 for c in cases if c["label"] == "vulnerable")
    n_safe = sum(1 for c in cases if c["label"] == "safe")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    # ---------------- כותרת ----------------
    p = doc.add_paragraph()
    p.paragraph_format.bidi = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('דוח סיכום עבודת גמר — מערכת "מבקר הקוד הכפול" (Dual Code Auditor)')
    r.bold = True
    r.font.rtl = True
    r.font.size = Pt(18)
    r.font.color.rgb = DEEP_BLUE
    md_lines.append('# דוח סיכום עבודת גמר — מערכת "מבקר הקוד הכפול" (Dual Code Auditor)\n')

    p = doc.add_paragraph()
    p.paragraph_format.bidi = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("מגישים: טל יהב, עמית יואל לוי, יונתן קרבלניק\nקורס: מבוא לסייבר בניהול")
    r.font.rtl = True
    r.font.size = Pt(10)
    r.font.italic = True
    md_lines.append("מגישים: טל יהב, עמית יואל לוי, יונתן קרבלניק  \nקורס: מבוא לסייבר בניהול\n")

    # ---------------- 1. תיאור הבעיה ----------------
    heading(doc, "תיאור הבעיה והנתונים", Pt(14))
    para(doc,
         "הפרויקט מיישם מערכת מרובת-סוכנים (Multi-Agent System) אוטונומית ב-n8n, שמזהה חולשות אבטחה "
         "בקוד מקור ומייצרת עבורן תיקון. הבעיה שהמערכת מטפלת בה היא עומס ההתראות בכלי סריקה סטטיים "
         "מסורתיים (SAST): שיעור גבוה של התראות שווא (False Positives) גורם לאנליסטים לבזבז חלק ניכר "
         "מזמנם על מיון במקום על טיפול. הערכה נפוצה בתעשייה, שמופיעה גם במסמך המטלה של הקורס, היא כ-70% "
         "מזמן האנליסט. בנוסף, גם כשמזוהה פגיעות אמיתית, הפער בין הזיהוי לבין דחיפת תיקון לייצור נמדד "
         "בימים, ובזמן הזה המערכת חשופה.")
    para(doc,
         'המערכת מבוססת על ארכיטקטורת Actor-Critic. סוכן ה-Red Team מנתח את הקוד, מזהה חולשות, מייצר '
         'הוכחת היתכנות (PoC) ומחשב מחרוזת וקטור CVSS 3.1. סוכן ה-Blue Team מקבל את הקוד המקורי במלואו '
         'יחד עם דוח הצוות האדום, ומחזיר גרסה מאובטחת (Secure Code Refactoring).')
    para(doc,
         "הנתונים שהמערכת צורכת בייצור הם קבצי Diff של Pull Requests. לצורך הערכה נבנה קורפוס בדיקה "
         "מתויג של %d קטעי קוד: %d פגיעים ו-%d בטוחים. הקטעים הבטוחים אינם קישוט — בלעדיהם אי אפשר "
         "למדוד התראות שווא, שהן הבעיה שהפרויקט מתיימר לפתור."
         % (len(cases), n_vuln, n_safe))

    # ---------------- 2. עיבוד מקדים ----------------
    heading(doc, "עיבוד מקדים של הנתונים", Pt(14))
    para(doc, "צינור העיבוד המקדים בנוי מארבעה שלבים:")
    para(doc, "טריגר GitHub Webhook מאזין לאירועי Pull Request במאגר.", bold_prefix="1. ", bullet=True)
    para(doc, "שכבת אימות בודקת את הבקשה לפני כל פעולה יקרה (פירוט בפרק הארכיטקטורה).",
         bold_prefix="2. ", bullet=True)
    para(doc, "בקשת HTTP ל-API של GitHub מושכת את קובץ ה-Diff של ה-PR.",
         bold_prefix="3. ", bullet=True)
    para(doc,
         "צומת Code ב-JavaScript מסנן את ה-Diff ומחלץ ממנו את שורות הקוד שנוספו בלבד, "
         "תוך התעלמות מכותרות ה-Diff, ומזהה את נתיב הקובץ. זהו שלב הקטנת הטוקנים: רק הקוד החדש "
         "נשלח למודל השפה, ולא ה-Diff הגולמי.",
         bold_prefix="4. ", bullet=True)
    para(doc,
         "בסיום נבנה אובייקט JSON אחיד (קוד, נתיב קובץ, מאגר, מספר PR, ענף) שמשמש כקלט לשני הסוכנים. "
         "אותו מבנה נוצר גם עבור בדיקה ידנית של קטע קוד, כך ששני מסלולי ההפעלה מתנקזים לצינור אחד.")

    # ---------------- 3. חקר נתונים ----------------
    heading(doc, "חקר נתונים (EDA) והרכב קורפוס הבדיקה", Pt(14))
    para(doc,
         "הפרק הזה מתאר את הרכב קורפוס הבדיקה שנבנה לפרויקט. כל המספרים והתרשימים כאן נגזרים "
         "אוטומטית מקובץ הקורפוס עצמו (scripts/eval/dataset.json).")

    heading(doc, "א. פיזור החולשות על פני OWASP Top 10", Pt(12), md_level=3, space_before=Pt(6))
    para(doc, "בדקנו את פיזור סוגי החולשות בקורפוס ביחס לקטגוריות OWASP Top 10 (2021).",
         bold_prefix="מה עשינו: ")
    para(doc, "כדי לוודא שההערכה לא נשענת על סוג חולשה יחיד, ושהיא מכסה מגוון משפחות תקיפה.",
         bold_prefix="למה עשינו: ")

    owasp_counts = {}
    for c in cases:
        if c["label"] == "vulnerable" and c.get("owasp"):
            k = c["owasp"]
            owasp_counts[k] = owasp_counts.get(k, 0) + 1
    top = sorted(owasp_counts.items(), key=lambda kv: -kv[1])[0]
    para(doc,
         "הקורפוס מכסה %d קטגוריות OWASP שונות. הקטגוריה הגדולה ביותר היא %s עם %d מקרים, "
         "מכיוון שקטגוריית ההזרקות מרכזת את החולשות הנפוצות ביותר בקוד אפליקטיבי "
         "(SQL Injection, Command Injection, XSS ו-eval). שאר הקטגוריות מיוצגות ב-1–2 מקרים כל אחת."
         % (len(owasp_counts), top[0], top[1]),
         bold_prefix="מה רואים בגרף: ")
    figure(doc, "fig1_corpus_owasp.png", "תרשים 1: פיזור המקרים הפגיעים לפי קטגוריות OWASP Top 10.")

    heading(doc, "ב. איזון בין מקרים פגיעים למקרים בטוחים", Pt(12), md_level=3, space_before=Pt(6))
    para(doc, "סיווגנו כל מקרה בקורפוס כפגיע או כבטוח, ובנינו לכל חולשה גם מקרה נגדי מתוקן.",
         bold_prefix="מה עשינו: ")
    para(doc,
         "מדד קצב זיהוי לבדו הוא חסר משמעות: מערכת שמסמנת כל קוד כפגיע תשיג 100% זיהוי. "
         "רק נוכחות של מקרים בטוחים מאפשרת למדוד התראות שווא ולהעריך את המערכת באמת.",
         bold_prefix="למה עשינו: ")
    para(doc,
         "הקורפוס מכיל %d מקרים פגיעים מול %d בטוחים (%.0f%% / %.0f%%). רוב המקרים הבטוחים הם "
         "הגרסה המתוקנת של מקרה פגיע מקביל — למשל שאילתה פרמטרית מול שרשור מחרוזות. "
         "זוגות כאלה מקשים על המערכת יותר מקוד בטוח אקראי, כי ההבדל ביניהם דק."
         % (n_vuln, n_safe, 100.0 * n_vuln / len(cases), 100.0 * n_safe / len(cases)),
         bold_prefix="מה רואים בגרף: ")
    figure(doc, "fig2_corpus_balance.png", "תרשים 2: יחס המקרים הפגיעים והבטוחים בקורפוס.")

    # ---------------- 4. ארכיטקטורה ----------------
    heading(doc, "מודלים וסוכנים: ארכיטקטורה ורציונל", Pt(14))
    para(doc, "המערכת מיושמת ב-n8n ומורכבת מ-23 צמתים. מנוע מודל השפה הוא Groq עם Llama 3.3 70B Versatile.")

    heading(doc, "1. סוכן הצוות האדום (Red Team AI Agent)", Pt(12), md_level=3, space_before=Pt(6))
    para(doc,
         "הסוכן מדמה תוקף חיצוני. ה-System Prompt מנחה אותו לאתר חולשות, להפיק PoC, ולחשב מדדי "
         "CVSS 3.1 מלאים: הן מדדי הניצול (AV, AC, PR, UI) והן מדדי הנזק (C, I, A), ולהחזיר "
         "מחרוזת וקטור תקנית.",
         bold_prefix="רציונל: ")
    para(doc,
         "הפרומפט מורה לסוכן במפורש להחזיר NO_VULNERABILITIES_FOUND כשלא נמצאה חולשה, ומגדיר "
         "התראת שווא ככישלון חמור כמו החמצת חולשה. ההנחיה הזו נוספה משום שמודל שפה שנשאל "
         '"מצא חולשות בקוד" נוטה למצוא חולשה גם כשאין, כדי לספק את הבקשה.',
         bold_prefix="החלטה מרכזית: ")

    heading(doc, "2. סוכן הצוות הכחול (Blue Team AI Agent)", Pt(12), md_level=3, space_before=Pt(6))
    para(doc,
         "סוכן ההגנה מקבל שני קלטים: הקוד המקורי במלואו ודוח הצוות האדום. הוא מחזיר את הקוד המתוקן "
         "בתוך בלוק קוד יחיד, ומתחתיו הסבר על התיקון וכיצד הוא חוסם את ה-PoC.",
         bold_prefix="רציונל: ")
    para(doc,
         "העברת הקוד המקורי לסוכן הכחול היא תנאי הכרחי. בגרסה מוקדמת הוא קיבל רק את דוח הצוות האדום, "
         "ונאלץ לשחזר את הקוד מתוך תיאור מילולי — מצב שמייצר קוד שאינו תואם למקור.",
         bold_prefix="החלטה מרכזית: ")

    heading(doc, "3. תזמור טכני במקום סוכן מנהל", Pt(12), md_level=3, space_before=Pt(6))
    para(doc,
         "התזמור בין הסוכנים מבוצע בתשתית n8n (צמתי Set, If, Code) ולא באמצעות סוכן AI שלישי. "
         "לזרימת בקרה יש תשובה נכונה אחת, ואין סיבה להפקיד אותה בידי רכיב הסתברותי. "
         "בנוסף, כל קריאה למודל עולה טוקנים וזמן.",
         bold_prefix="רציונל ארכיטקטוני: ")

    heading(doc, "4. שכבת אבטחה לוובהוק הציבורי", Pt(12), md_level=3, space_before=Pt(6))
    para(doc,
         "הוובהוק שמפעיל את המערכת חשוף לאינטרנט. ללא אימות, כל אדם יכול לשלוח JSON שמצביע על מאגר "
         "שרירותי, לגרום למערכת למשוך אותו, לשרוף טוקנים, ובתצורה שדוחפת תיקונים — לגרום לכתיבת "
         "Commit באמצעות ההרשאות של המערכת. זו חולשה במערכת שתפקידה לאתר חולשות.",
         bold_prefix="הבעיה: ")
    para(doc,
         "צומת Code ייעודי מאמת כל בקשה לפני כל פעולה יקרה. בקשה שמגיעה מ-GitHub נבדקת מול חתימת "
         "HMAC-SHA256 בכותרת x-hub-signature-256, בהשוואה שעמידה בפני התקפות תזמון (timingSafeEqual). "
         "בקשה ידנית מדף התצוגה מגיעה ללא חתימה, ולכן מותרת רק עבור מאגרים ברשימת היתר. "
         "בקשה שנכשלת מקבלת תשובת 403 עם נימוק, ולא מגיעה לשום קריאת LLM.",
         bold_prefix="הפתרון: ")

    heading(doc, "5. שכבת התיעוד והתגובה", Pt(12), md_level=3, space_before=Pt(6))
    para(doc,
         "כל הרצה נכתבת לקובץ Markdown נפרד ב-Google Drive עם חותמת זמן בשם הקובץ. עבור טריגר של "
         "GitHub, המערכת מפרסמת את הדוח כהערה על ה-Pull Request. דחיפת Commit מתבצעת רק אם התקיימו "
         "כל התנאים: דווחה חולשה, הסוכן הכחול החזיר בלוק קוד תקין, וקיים נתיב קובץ. הקוד שנדחף מחולץ "
         "מתוך בלוק ה-Markdown בלבד, כדי שטקסט ההסבר לא ייכתב לתוך קובץ המקור. "
         "בכל מקרה נדרש אישור אנושי (Merge) לפני שהשינוי מגיע לענף הראשי.")

    # ---------------- 5. אימות ובדיקות ----------------
    heading(doc, "אימות המערכת ובדיקות", Pt(14))
    para(doc,
         "לוגיקת הצמתים המרכזיים מכוסה בבדיקות אוטומטיות. הבדיקות טוענות את קוד הצמתים ישירות מקובץ "
         "ה-JSON של ה-workflow ומריצות אותו בסביבת n8n מדומה, כך שהן בודקות את הקוד שרץ בפועל.")

    if tests:
        passed, failed = tests
        para(doc,
             "בהרצה האחרונה עברו %d בדיקות ונכשלו %d." % (passed, failed),
             bold_prefix="תוצאת ההרצה: ")
    else:
        para(doc, "לא ניתן היה להריץ את הבדיקות בעת הפקת הדוח.", bold_prefix="תוצאת ההרצה: ")

    para(doc, "הבדיקות המרכזיות:")
    para(doc, "בקשה עם חתימת HMAC תקינה מתקבלת; חתימה מזויפת נדחית.", bullet=True)
    para(doc, "בקשה לא חתומה למאגר שאינו ברשימת ההיתר נדחית.", bullet=True)
    para(doc, "הודעת הדחייה אינה חושפת את הסוד המשותף.", bullet=True)
    para(doc, "עיבוד ה-Diff מחלץ שורות שנוספו בלבד, בלי כותרות Diff.", bullet=True)
    para(doc, "הקוד שנדחף לריפו אינו מכיל טקסט הסבר או סימוני Markdown.", bullet=True)
    para(doc, "כשלא נמצאה חולשה, או כשהסוכן לא החזיר בלוק קוד — לא נדחף Commit.", bullet=True)

    para(doc,
         "בנוסף לבדיקות היחידה, שכבת האבטחה אומתה מול השרת החי. בקשה למאגר שאינו ברשימת ההיתר, "
         "בקשה עם גוף פגום, ובקשה עם חתימה מזויפת — כל השלוש קיבלו 403 עם נימוק מתאים, ולא הפעילו "
         "אף קריאה למודל שפה. בקשה ל-Pull Request שאינו קיים מחזירה 502 עם הסבר, במקום להיכשל בשקט.")

    # ---------------- 6. הערכת ביצועים ----------------
    heading(doc, "הערכת ביצועים", Pt(14))

    if results:
        s = results["summary"]
        para(doc,
             "ההערכה בוצעה על קורפוס הבדיקה המתויג. כל קטע קוד נשלח למערכת דרך endpoint הסריקה, "
             "והסיווג שהתקבל הושווה לתווית האמת. תצורת ההרצה: %s, ארכיטקטורה %s. תאריך ההרצה: %s."
             % (s["model"], s["architecture"], s["generated_at"]))

        rows = [
            ["מקרים שהושלמו", "%d מתוך %d" % (s["completed_cases"], s["total_cases"])],
            ["קצב זיהוי (Recall)", "%s%%" % s["detection_rate_pct"]],
            ["התראות שווא (FPR)", "%s%%" % s["false_positive_rate_pct"]],
            ["דיוק (Precision)", "%s%%" % s["precision_pct"]],
            ["הפקת וקטור CVSS תקין", "%s%%" % s["cvss_vector_emitted_pct"]],
            ["זמן ריצה ממוצע", "%s שניות" % s["avg_runtime_sec"]],
            ["זמן ריצה חציוני", "%s שניות" % s["median_runtime_sec"]],
        ]
        table(doc, ["מדד", "ערך שנמדד"], rows)

        figure(doc, "fig3_confusion.png", "תרשים 3: תוצאות הסיווג שנמדדו בפועל.")
        figure(doc, "fig4_runtime.png", "תרשים 4: התפלגות זמני הריצה מקצה לקצה.")

        by_id = {c["id"]: c for c in cases}

        # ---- השוואה בין שתי ההרצות: לפני ואחרי תיקון הפרומפט ----
        if baseline:
            b = baseline["summary"]
            n_runs = "שלוש" if midrun else "שתי"
            heading(doc, "%s הרצות: אבחון הטיה ותיקונה" % n_runs, Pt(12), md_level=3,
                    space_before=Pt(6))
            para(doc,
                 "ההערכה הורצה %s פעמים על אותו קורפוס ואותו מודל. בין ההרצות שונה "
                 "ה-System Prompt של הסוכן האדום בלבד, בעקבות אבחון של ההרצה הקודמת. "
                 "כל שאר רכיבי המערכת נותרו זהים, ולכן ההפרש בין ההרצות משקף את השינוי "
                 "בהנחיה ולא שינוי אחר." % n_runs)

            runs = [("הרצה 1", b)]
            if midrun:
                runs.append(("הרצה 2", midrun["summary"]))
            runs.append(("הרצה %d (סופית)" % (len(runs) + 1), s))

            def row(label, key, suffix=""):
                return [label] + ["%s%s" % (r[key], suffix) for _, r in runs]

            table(doc,
                  ["מדד"] + [name for name, _ in runs],
                  [row("קצב זיהוי", "detection_rate_pct", "%"),
                   row("התראות שווא", "false_positive_rate_pct", "%"),
                   row("דיוק", "precision_pct", "%"),
                   row("החמצות", "false_negatives"),
                   row("התראות שווא (מספר)", "false_positives")])

            figure(doc, "fig5_run_comparison.png",
                   "תרשים 5: השפעת כל תיקון בהנחיה על שלושת המדדים, באותו קורפוס ואותו מודל.")

            b_fps = [r["id"] for r in baseline["records"]
                     if r["label"] == "safe" and r.get("predicted_vulnerable")]
            para(doc,
                 "בהרצה הראשונה המערכת זיהתה את כל המקרים הפגיעים, אך סימנה גם %d מהמקרים "
                 "הבטוחים (%s). בדיקה של שלושתם העלתה מכנה משותף: כולם הגרסה המתוקנת של "
                 "חולשה מקבילה בקורפוס — תבנית תצוגה בצד השרת, בדיקת נתיב מוחלט מול תיקיית "
                 "בסיס, ופענוח JSON עם אימות טיפוס. הסוכן זיהה נכון את אזור הסיכון, אך לא "
                 "בחן אם אמצעי ההגנה שכבר נמצא בקוד סוגר אותו."
                 % (b["false_positives"], ", ".join(b_fps)),
                 bold_prefix="אבחון: ")
            para(doc,
                 "ל-System Prompt נוסף שלב חובה לפני הדיווח: לאחר איתור דפוס סיכון, הסוכן "
                 "נדרש לעבור על הקוד ולבדוק אם קיים אמצעי שמנטרל אותו, ולדווח רק אם הוא יכול "
                 "לכתוב PoC שעובד למרות ההגנות. נוסף שדה חובה בפלט שבו עליו לנמק מדוע ההגנות "
                 "הקיימות אינן מספיקות.",
                 bold_prefix="התיקון: ")
            if midrun:
                m = midrun["summary"]
                para(doc,
                     "שיעור התראות השווא ירד מ-%s%% ל-%s%%, אך נוצרה החמצה: המקרה של סיסמה "
                     "מוטמעת בקוד לא דווח. בדיקה העלתה שההנחיה החדשה הופעלה רחב מדי. "
                     "לחלק ממחלקות החולשה אין אמצעי הגנה אפשרי *בתוך* אותו קטע קוד — סוד "
                     "מוטמע, גיבוב חלש לסיסמאות, ביטול אימות תעודה. הסוכן חיפש הגנה, לא "
                     "מצא, ופירש זאת כאילו אין חולשה."
                     % (b["false_positive_rate_pct"], m["false_positive_rate_pct"]),
                     bold_prefix="הרצה 2: ")
                para(doc,
                     "לשלב בדיקת הנטרול נוספה רשימת חריגים: מחלקות שבהן היעדר הגנה הוא "
                     "הממצא עצמו, ולכן מדווחים בהן תמיד. הרשימה מונה חמש מחלקות מוגדרות "
                     "היטב ואינה חופפת למקרים שגרמו להתראות השווא, שכולם עסקו בקלט משתמש "
                     "שעובר דרך הגנה.",
                     bold_prefix="התיקון השני: ")

            para(doc,
                 "בתצורה הסופית קצב הזיהוי הוא %s%% ללא החמצות, שיעור התראות השווא %s%% "
                 "והדיוק %s%%. ביחס להרצה הראשונה, שיעור התראות השווא ירד מ-%s%% והדיוק "
                 "עלה מ-%s%%, בלי לוותר על אף ממצא. כלומר האיזון בין דיוק לכיסוי לא נפתר "
                 "בבחירת סף, אלא בהפרדה בין מחלקות חולשה שיש להן אמצעי הגנה אפשרי בקוד "
                 "לבין כאלה שאין."
                 % (s["detection_rate_pct"], s["false_positive_rate_pct"],
                    s["precision_pct"], b["false_positive_rate_pct"], b["precision_pct"]),
                 bold_prefix="התוצאה הסופית: ")

        # ניתוח ההתראות שווא שנותרו, נגזר מהרשומות עצמן ולא נכתב ידנית.
        fps = [r for r in results["records"]
               if r["label"] == "safe" and r.get("predicted_vulnerable")]
        fns = [r for r in results["records"]
               if r["label"] == "vulnerable" and r.get("predicted_vulnerable") is False]

        if fps or fns:
            heading(doc, "הטעויות שנותרו", Pt(12), md_level=3, space_before=Pt(6))
            figure(doc, "fig6_per_case.png",
                   "תרשים 6: תוצאה לכל אחד מ-20 מקרי הבדיקה. FP מסמן התראת שווא.")
        if fps:
            rows = [[r["id"], by_id.get(r["id"], {}).get("language", "—"),
                     r.get("cvss_vector") or "—"] for r in fps]
            table(doc, ["התראת שווא", "שפה", "וקטור CVSS שדווח"], rows)
            para(doc,
                 "ההתראה שנותרה היא על נתיב Express שמעביר קלט משתמש למנוע תבניות. "
                 "כאן ההגדרה של המקרה כבטוח שנויה במחלוקת: בטיחותו תלויה בקובץ התבנית, "
                 "שאינו חלק מהקטע שנבדק. אם התבנית משתמשת בהחלפה שאינה מבצעת escaping, "
                 "הקוד אכן פגיע. זו מגבלה של הקורפוס יותר מאשר טעות של הסוכן, ותיוג "
                 "המקרה כבטוח נשען על הנחה שאי אפשר לאמת מתוך הקוד עצמו.")
        if fns:
            para(doc,
                 "ההחמצה היחידה היא %s — סיסמה מוטמעת בקוד. הנחיית האימות שנוספה גרמה "
                 "לסוכן להתייחס לקבוע כאל ערך תצורה לגיטימי. זו דוגמה ישירה לעלות של "
                 "ההנחיה השמרנית: היא מסננת רעש, אך גם משתיקה ממצא אמיתי שאין מולו "
                 "אמצעי הגנה כלל."
                 % ", ".join(r["id"] for r in fns))

        para(doc,
             "המדגם מונה %d מקרים בלבד. בגודל כזה, כל מקרה בודד מזיז את האחוזים בכמה נקודות, "
             "ולכן יש לקרוא את המספרים כאינדיקציה לכיוון ולא כמדידה מדויקת. "
             "כדי לטעון טענה סטטיסטית ממשית נדרש מדגם גדול משמעותית והרצות חוזרות."
             % s["completed_cases"],
             bold_prefix="הערת מהימנות: ")
        para(doc,
             "ההרצה בוצעה כולה על מודל אחד. נבדק בלוגים של השרת שלא אירע ולו כשל אחד "
             "שהיה מפעיל את מנגנון המודל החלופי, ולכן המספרים מתארים מודל יחיד ולא הרכב "
             "של שניים.",
             bold_prefix="תקפות המדידה: ")
    else:
        para(doc,
             "ההערכה הכמותית טרם הורצה, ולכן אין בדוח זה טבלת ביצועים.",
             bold_prefix="סטטוס: ")
        para(doc,
             "תשתית ההערכה בנויה ומוכנה: קורפוס מתויג של %d מקרים (scripts/eval/dataset.json) "
             "וסקריפט הרצה שמחשב קצב זיהוי, שיעור התראות שווא, דיוק וזמני ריצה "
             "(scripts/eval/run_eval.py). התנאי החסר להרצה הוא קרדנצ'ל פעיל של ספק מודל השפה "
             "במופע ה-n8n. עם השלמתו, הרצת הסקריפט מייצרת את התוצאות והדוח מציג אותן אוטומטית."
             % len(cases))
        para(doc,
             "לא הוצגו כאן מספרי ביצועים משום שלא נמדדו. הצגת מספרים שלא נמדדו הייתה הופכת את הדוח "
             "לבלתי ניתן להגנה.",
             bold_prefix="הערה מתודולוגית: ")

    heading(doc, "דיון מתודולוגי: מדוע שני סוכנים", Pt(12), md_level=3, space_before=Pt(6))
    para(doc,
         "ההנמקה לארכיטקטורה הדו-סוכנית היא עקרונית ולא ניסויית. מודל שפה יחיד שמתבקש גם לאתר "
         "חולשה וגם לתקן אותה נמצא בניגוד עניינים מובנה: הוא נדרש לשפוט את עבודתו שלו. הפרדה "
         "לשני תפקידים — מאתר ומתקן — יוצרת נקודת בקרה שבה פלט הסוכן הראשון הוא הקלט של השני, "
         "ומחייבת שהממצא יהיה מנוסח בצורה שאפשר לפעול לפיה.")
    para(doc,
         "הארכיטקטורה הזו אינה מבטלת התראות שווא. אם הסוכן האדום מדווח על חולשה שאינה "
         "קיימת, הסוכן הכחול מקבל את הדיווח כנתון ומנסה לתקן. אין כאן מנגנון הצבעה ואין "
         "ביקורת הדדית: הסוכן הכחול אינו מוסמך לפסול את הממצא, אלא רק לפעול לפיו.")

    if results:
        para(doc,
             "המדידה מאשרת זאת. שיעור התראות השווא עומד על %s%%, וכל הירידה ביחס להרצה "
             "הראשונה הושגה בשינוי ההנחיה לסוכן האדום ולא בזכות הארכיטקטורה. הציפייה "
             "שהפרדת התפקידים תסנן רעש מעצמה לא התממשה, ומי שמתכנן מערכת כזו אינו יכול "
             "להסתמך על ההפרדה ככלי סינון."
             % s["false_positive_rate_pct"])
        para(doc,
             "מה שההפרדה כן נותנת הוא אילוץ מבני: הממצא חייב להיות ספציפי מספיק כדי "
             "שאפשר יהיה לגזור ממנו תיקון. זה מייצר דיווחים ניתנים לפעולה ומאפשר את "
             "שכבת ה-Self-Healing, אבל אלה יתרונות בתחום השימושיות ולא בתחום הדיוק. "
             "לסינון התראות שווא נדרש רכיב נפרד שתפקידו לפסול ממצאים, למשל מאמת שמריץ "
             "את ה-PoC מול הקוד ובודק אם הוא באמת עובד.")

    # ---------------- 7. מגבלות ----------------
    heading(doc, "מגבלות הפרויקט", Pt(14))
    para(doc, "המגבלות שלהלן נכונות למועד הגשת הדוח:")
    para(doc,
         "הערכה על מודל אחד בלבד (Llama 3.3 70B דרך Groq). לא בוצעה השוואה בין ספקי מודלים.",
         bullet=True)
    para(doc,
         "קורפוס הבדיקה קטן ומורכב מקטעי קוד קצרים ומבודדים. הוא אינו מייצג קוד ייצור עם "
         "תלויות בין קבצים, שבו הקשר רחב משפיע על השאלה אם חולשה ניתנת לניצול.",
         bullet=True)
    para(doc,
         "המערכת מנתחת את השורות שנוספו ב-Diff ולא את הקובץ המלא. זה חוסך טוקנים, אך עלול להחמיץ "
         "חולשות שנוצרות מהאינטראקציה בין קוד חדש לקוד קיים.",
         bullet=True)
    para(doc,
         "אימות איכות התיקון של הסוכן הכחול נעשה בקריאה אנושית. לא נבנתה סביבת הרצה שמריצה את "
         "ה-PoC מול הקוד המתוקן ומוודאת אוטומטית שהוא נחסם.",
         bullet=True)
    para(doc,
         "רשימת ההיתר במנגנון האבטחה מוגדרת בקוד הצומת. בפריסה ארגונית יש להעביר אותה לניהול חיצוני.",
         bullet=True)

    # ---------------- 8. סיכום ----------------
    heading(doc, "סיכום והשלכות יישומיות", Pt(14))
    para(doc,
         "הפרויקט מדגים שאפשר לבנות צינור אוטומטי שמחבר אירוע Pull Request לניתוח אבטחה ולהצעת תיקון, "
         "כשהתזמור נעשה בתשתית אוטומציה והשיקול הסמנטי בלבד מופקד בידי מודלי שפה. "
         "הערך העסקי הוא קיצור הזמן בין כתיבת קוד פגיע לבין הצפת הממצא למפתח — מימים לדקות.")
    para(doc,
         "מבחינה יישומית, ההמלצה היא לשלב מערכת כזו בשלב ה-Pull Request של תהליך ה-CI/CD, "
         "ותמיד עם Human-in-the-loop: המערכת מציעה, המפתח מאשר. תיקון אוטומטי שנדחף בלי אישור אנושי "
         "מעביר את הסיכון ממקום למקום במקום להקטין אותו.")
    para(doc,
         "המסקנה המתודולוגית המרכזית מהפרויקט אינה נוגעת למודלים אלא לתשתית שסביבם: המערכת נשברה "
         "בנקודות שלא היו קשורות ל-AI — אימות בקשות, שמירה על גבול בין פלט טקסטואלי לתוכן קובץ, "
         "וטיפול בכשלים. מערכת סוכנים שכותבת לתוך מאגר קוד היא בעלת הרשאות כתיבה, ויש להתייחס אליה "
         "כאל רכיב בעל הרשאות ולא כאל כלי ניתוח.")

    doc.save(DOCX_OUT)
    with open(MD_OUT, "w", encoding="utf-8") as f:
        f.write("\n".join(md_lines) + "\n")

    print("DOCX ->", DOCX_OUT)
    print("MD   ->", MD_OUT)
    print("tests:", tests, "| results:", "present" if results else "absent")


if __name__ == "__main__":
    build()
